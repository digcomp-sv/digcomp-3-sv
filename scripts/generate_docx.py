#!/usr/bin/env python3
"""
generate_docx.py
----------------
Genererar ett remissklart Word-dokument (.docx) från JSON-datan i `data/sv/`.

Syfte: göra det möjligt att släppa läsbara versioner av översättningen för
remiss och kommentering utan att manuellt underhålla ett parallellt
dokument. Datan i JSON-filerna är alltid source of truth.

Användning:
    python scripts/generate_docx.py [--output FIL] [--status-filter STATUS]

Exempel:
    # Generera hela dokumentet med all data
    python scripts/generate_docx.py

    # Bara posterna som är översatta (draft eller senare)
    python scripts/generate_docx.py --status-filter draft

    # Ange utdatafil
    python scripts/generate_docx.py --output releases/v0.1-draft/digcomp-3-sv.docx

Beroenden:
    pip install python-docx
"""
from __future__ import annotations

import argparse
import json
import sys
from datetime import date
from pathlib import Path

try:
    from docx import Document
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.stderr.write(
        "python-docx är inte installerat. Kör: pip install python-docx\n"
    )
    sys.exit(1)


# ---------- Konstanter ----------

REPO_ROOT = Path(__file__).resolve().parent.parent
DATA_DIR = REPO_ROOT / "data" / "sv"
DEFAULT_OUTPUT = REPO_ROOT / "releases" / f"digcomp-3-sv-{date.today().isoformat()}.docx"

# Färger
COLOR_AREA_1 = RGBColor(0xF5, 0xB9, 0x3C)  # gul (info search) - matchar JRC
COLOR_AREA_2 = RGBColor(0x3F, 0x7A, 0xB3)  # blå (communication)
COLOR_AREA_3 = RGBColor(0xE8, 0x82, 0x3C)  # orange (content creation)
COLOR_AREA_4 = RGBColor(0x6F, 0xAA, 0x6E)  # grön (safety, wellbeing)
COLOR_AREA_5 = RGBColor(0xD8, 0x6B, 0x5E)  # röd (problem solving)

AREA_COLORS = {1: COLOR_AREA_1, 2: COLOR_AREA_2, 3: COLOR_AREA_3, 4: COLOR_AREA_4, 5: COLOR_AREA_5}

COLOR_MUTED = RGBColor(0x66, 0x66, 0x66)
COLOR_PLACEHOLDER = RGBColor(0xAA, 0x22, 0x22)

STATUS_ORDER = ["not_started", "draft", "needs_review", "reviewed", "final"]

STATUS_LABELS = {
    "not_started": "Ej påbörjad",
    "draft": "Utkast",
    "needs_review": "Behöver granskning",
    "reviewed": "Granskad",
    "final": "Slutlig",
}


# ---------- Datamodell ----------

def load_json(name: str) -> list | dict:
    path = DATA_DIR / name
    if not path.exists():
        sys.exit(f"Hittar inte {path}. Kör scriptet från repo-roten eller kontrollera data/sv/.")
    with path.open(encoding="utf-8") as f:
        return json.load(f)


def load_all_data() -> dict:
    return {
        "areas": load_json("competence_areas.json"),
        "competences": load_json("competences.json"),
        "levels": load_json("proficiency_levels.json"),
        "statements": load_json("competence_statements.json"),
        "outcomes": load_json("learning_outcomes.json"),
        "glossary": load_json("glossary.json"),
    }


def status_at_least(item_status: str, threshold: str) -> bool:
    """Returnerar True om postens status är minst så långt som threshold."""
    try:
        return STATUS_ORDER.index(item_status) >= STATUS_ORDER.index(threshold)
    except ValueError:
        return False


# ---------- Dokumentbyggande ----------

def add_page_break(doc: Document) -> None:
    doc.add_paragraph().add_run().add_break()


def set_cell_background(cell, color_hex: str) -> None:
    """Sätter bakgrundsfärg på en tabellcell (finns inte i python-docx API)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def render_text_pair(paragraph, en_text: str, sv_text: str | None, show_en: bool = True) -> None:
    """Renderar svensk text (med fallback till markerad engelsk om ej översatt)."""
    if sv_text:
        run = paragraph.add_run(sv_text)
    else:
        run = paragraph.add_run(f"[Ej översatt: {en_text}]")
        run.italic = True
        run.font.color.rgb = COLOR_PLACEHOLDER

    if show_en and sv_text:
        # Lägg till engelsk referens i muted färg
        paragraph.add_run("  ")
        en_run = paragraph.add_run(f"[EN: {en_text}]")
        en_run.italic = True
        en_run.font.size = Pt(9)
        en_run.font.color.rgb = COLOR_MUTED


def render_status_badge(paragraph, status: str) -> None:
    label = STATUS_LABELS.get(status, status)
    run = paragraph.add_run(f"  [{label}]")
    run.italic = True
    run.font.size = Pt(8)
    run.font.color.rgb = COLOR_MUTED


def render_ai_label(paragraph, ai_label: str | None) -> None:
    if not ai_label or ai_label == "AI not Implicit or Explicit":
        return
    run = paragraph.add_run(f" [{ai_label}]")
    run.font.size = Pt(8)
    run.bold = True
    if ai_label == "AI-Explicit":
        run.font.color.rgb = RGBColor(0x80, 0x00, 0x80)
    else:
        run.font.color.rgb = COLOR_MUTED


# ---------- Sektioner ----------

def build_title_page(doc: Document, stats: dict) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("DigComp 3.0")
    run.bold = True
    run.font.size = Pt(36)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Europeiskt ramverk för medborgares digitala kompetens")
    run.font.size = Pt(18)

    version = doc.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = version.add_run("Svensk översättning – remissversion")
    run.italic = True
    run.font.size = Pt(14)

    doc.add_paragraph()
    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Genererad: {date.today().isoformat()}\n")
    meta.add_run(f"Innehåll: {stats['areas']} områden, {stats['competences']} kompetenser, ")
    meta.add_run(f"{stats['statements']} kompetenspåståenden, {stats['outcomes']} learning outcomes, ")
    meta.add_run(f"{stats['glossary']} termer\n")
    meta.add_run(f"Översättningsgrad: {stats['translation_percent']:.1f} %")

    doc.add_paragraph()

    source_p = doc.add_paragraph()
    source_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = source_p.add_run("Baserad på: Cosgrove, J. & Cachia, R. (2025).\n")
    run.font.size = Pt(10)
    run = source_p.add_run("DigComp 3.0: European Digital Competence Framework – Fifth Edition.\n")
    run.italic = True
    run.font.size = Pt(10)
    run = source_p.add_run("Publications Office of the European Union, Luxembourg. JRC144121.")
    run.font.size = Pt(10)

    doc.add_paragraph()
    licens = doc.add_paragraph()
    licens.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = licens.add_run("CC BY-SA 4.0")
    run.font.size = Pt(10)

    add_page_break(doc)


def build_toc_section(doc: Document, data: dict) -> None:
    doc.add_heading("Innehåll", level=1)

    # Kompetensöversikt som tabell
    for area in data["areas"]:
        h = doc.add_heading(level=2)
        sv_name = area["name"].get("sv")
        en_name = area["name"]["en"]
        run = h.add_run(f"{area['id']}. {sv_name or en_name}")
        if not sv_name:
            run.italic = True
            run.font.color.rgb = COLOR_PLACEHOLDER

        # Beskrivning
        p = doc.add_paragraph()
        render_text_pair(p, area["descriptor"]["en"], area["descriptor"].get("sv"), show_en=False)

        # Lista kompetenser i detta område
        area_competences = [c for c in data["competences"] if c["area_id"] == area["id"]]
        for comp in area_competences:
            list_p = doc.add_paragraph(style="List Bullet")
            sv_name = comp["name"].get("sv")
            en_name = comp["name"]["en"]
            run = list_p.add_run(f"{comp['id']} ")
            run.bold = True
            if sv_name:
                list_p.add_run(sv_name)
            else:
                run2 = list_p.add_run(f"[Ej översatt: {en_name}]")
                run2.italic = True
                run2.font.color.rgb = COLOR_PLACEHOLDER

    add_page_break(doc)


def build_proficiency_levels_section(doc: Document, data: dict) -> None:
    doc.add_heading("Kompetensnivåer", level=1)

    p = doc.add_paragraph()
    p.add_run(
        "DigComp 3.0 använder fyra kompetensnivåer, som också kan mappas till "
        "åtta nivåer motsvarande EQF (European Qualifications Framework)."
    )

    doc.add_heading("Fyra nivåer", level=2)
    for level in data["levels"]["four_level"]:
        lvl_heading = doc.add_paragraph()
        run = lvl_heading.add_run(level["name"].get("sv") or level["name"]["en"])
        run.bold = True
        run.font.size = Pt(13)
        if not level["name"].get("sv"):
            run.font.color.rgb = COLOR_PLACEHOLDER

        desc_p = doc.add_paragraph()
        render_text_pair(desc_p, level["description"]["en"], level["description"].get("sv"), show_en=False)

        purpose_p = doc.add_paragraph()
        run = purpose_p.add_run("Syfte: ")
        run.italic = True
        render_text_pair(purpose_p, level["purpose"]["en"], level["purpose"].get("sv"), show_en=False)

    add_page_break(doc)


def build_competence_section(doc: Document, data: dict, status_filter: str | None) -> None:
    """Huvudsektion: en underrubrik per kompetens, med CS-påståenden och LO:n."""
    doc.add_heading("Kompetenser", level=1)

    competences_by_id = {c["id"]: c for c in data["competences"]}
    statements_by_comp = {}
    for s in data["statements"]:
        statements_by_comp.setdefault(s["competence_id"], []).append(s)
    outcomes_by_comp = {}
    for o in data["outcomes"]:
        outcomes_by_comp.setdefault(o["competence_id"], []).append(o)

    for area in data["areas"]:
        area_heading = doc.add_heading(level=2)
        sv_name = area["name"].get("sv")
        en_name = area["name"]["en"]
        run = area_heading.add_run(f"Område {area['id']}: {sv_name or en_name}")

        area_competences = [c for c in data["competences"] if c["area_id"] == area["id"]]
        for comp in area_competences:
            # Kompetensrubrik
            comp_heading = doc.add_heading(level=3)
            sv_name = comp["name"].get("sv")
            en_name = comp["name"]["en"]
            run = comp_heading.add_run(f"{comp['id']} ")
            run.bold = True
            if sv_name:
                comp_heading.add_run(sv_name)
            else:
                run2 = comp_heading.add_run(en_name)
                run2.italic = True
                run2.font.color.rgb = COLOR_PLACEHOLDER

            # Beskrivning
            desc_p = doc.add_paragraph()
            run = desc_p.add_run("Beskrivning: ")
            run.italic = True
            render_text_pair(desc_p, comp["descriptor"]["en"], comp["descriptor"].get("sv"), show_en=False)

            # --- Kompetenspåståenden (CS) grupperade per nivå ---
            doc.add_paragraph().add_run("Kompetenspåståenden").bold = True

            statements = statements_by_comp.get(comp["id"], [])
            # Filter
            if status_filter:
                statements = [s for s in statements if status_at_least(s["translation_status"], status_filter)]

            # Gruppera per nivå
            levels_in_order = ["Basic", "Intermediate", "Advanced", "Highly Advanced"]
            for level in levels_in_order:
                level_statements = [s for s in statements if s["proficiency_level"] == level]
                if not level_statements:
                    continue

                level_p = doc.add_paragraph()
                run = level_p.add_run(f"  {level}")
                run.bold = True
                run.italic = True

                for stmt in level_statements:
                    bullet = doc.add_paragraph(style="List Bullet")
                    run = bullet.add_run(f"{stmt['id']}: ")
                    run.bold = True
                    render_text_pair(bullet, stmt["text"]["en"], stmt["text"].get("sv"), show_en=False)
                    render_ai_label(bullet, stmt.get("ai_label"))
                    if stmt["translation_status"] not in ("final", "reviewed"):
                        render_status_badge(bullet, stmt["translation_status"])

            # --- Learning outcomes (LO) grupperade per nivå och dimension ---
            doc.add_paragraph().add_run("Learning outcomes").bold = True

            outcomes = outcomes_by_comp.get(comp["id"], [])
            if status_filter:
                outcomes = [o for o in outcomes if status_at_least(o["translation_status"], status_filter)]

            # Tabell: 4 kolumner - ID | Nivå | Dim | Text
            if outcomes:
                table = doc.add_table(rows=1, cols=4)
                table.style = "Light Grid Accent 1"
                hdr = table.rows[0].cells
                hdr[0].text = "ID"
                hdr[1].text = "Nivå"
                hdr[2].text = "Dim."
                hdr[3].text = "Text"
                for h in hdr:
                    for p in h.paragraphs:
                        for r in p.runs:
                            r.bold = True

                # Sortera outcomes enligt nivå + dim
                level_order = {"Basic": 0, "Intermediate": 1, "Advanced": 2, "Highly advanced": 3}
                dim_order = {"Knowledge": 0, "Skill": 1, "Attitude": 2}
                outcomes_sorted = sorted(
                    outcomes,
                    key=lambda o: (
                        level_order.get(o["proficiency_level"], 99),
                        dim_order.get(o["dimension"], 99),
                        o["id"],
                    ),
                )

                for o in outcomes_sorted:
                    row = table.add_row().cells
                    row[0].text = o["id"]
                    row[1].text = o["proficiency_level"][:4]  # Bas/Inte/Adva/High
                    row[2].text = o["dimension"][:4]  # Know/Skil/Atti
                    cell_p = row[3].paragraphs[0]
                    render_text_pair(cell_p, o["text"]["en"], o["text"].get("sv"), show_en=False)
                    render_ai_label(cell_p, o.get("ai_label"))

                    for c in row[:3]:
                        for p in c.paragraphs:
                            for r in p.runs:
                                r.font.size = Pt(9)

            doc.add_paragraph()  # luft mellan kompetenser

        add_page_break(doc)


def build_glossary_section(doc: Document, data: dict) -> None:
    doc.add_heading("Ordlista", level=1)

    p = doc.add_paragraph()
    p.add_run(
        "Termerna är ordnade alfabetiskt efter engelska termen. Svenska termer som ännu inte "
        "är fastställda är markerade med kursiverad röd text."
    )

    glossary_sorted = sorted(data["glossary"], key=lambda g: g["term"]["en"].lower())

    for entry in glossary_sorted:
        en_term = entry["term"]["en"]
        sv_term = entry["term"].get("sv")

        term_p = doc.add_paragraph()
        if sv_term:
            run = term_p.add_run(sv_term)
            run.bold = True
            run.font.size = Pt(12)
            run = term_p.add_run(f"  ({en_term})")
            run.italic = True
            run.font.color.rgb = COLOR_MUTED
        else:
            run = term_p.add_run(en_term)
            run.bold = True
            run.font.size = Pt(12)
            run = term_p.add_run("  [svensk term ej fastställd]")
            run.italic = True
            run.font.color.rgb = COLOR_PLACEHOLDER

        def_p = doc.add_paragraph()
        render_text_pair(def_p, entry["explanation"]["en"], entry["explanation"].get("sv"), show_en=False)

        if entry.get("swedish_source"):
            src_p = doc.add_paragraph()
            run = src_p.add_run(f"Källa för svensk term: {entry['swedish_source']}")
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = COLOR_MUTED


# ---------- Huvudflöde ----------

def compute_stats(data: dict) -> dict:
    all_translatable = [
        *data["statements"],
        *data["outcomes"],
    ]
    translated = sum(
        1 for item in all_translatable
        if status_at_least(item.get("translation_status", "not_started"), "draft")
    )
    total = len(all_translatable)
    return {
        "areas": len(data["areas"]),
        "competences": len(data["competences"]),
        "statements": len(data["statements"]),
        "outcomes": len(data["outcomes"]),
        "glossary": len(data["glossary"]),
        "translation_percent": (translated / total * 100) if total else 0,
    }


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Generera Word-dokument från DigComp-3-sv JSON-datan."
    )
    parser.add_argument(
        "--output", "-o",
        type=Path,
        default=DEFAULT_OUTPUT,
        help=f"Sökväg för utdatafilen. Standard: {DEFAULT_OUTPUT}",
    )
    parser.add_argument(
        "--status-filter",
        choices=STATUS_ORDER,
        default=None,
        help="Ta bara med poster med minst denna status. Standard: alla poster.",
    )
    args = parser.parse_args()

    print(f"Läser JSON-data från {DATA_DIR}...")
    data = load_all_data()
    stats = compute_stats(data)

    print(f"  {stats['areas']} områden, {stats['competences']} kompetenser")
    print(f"  {stats['statements']} CS, {stats['outcomes']} LO, {stats['glossary']} termer")
    print(f"  Översättningsgrad: {stats['translation_percent']:.1f} %")

    if args.status_filter:
        print(f"  Filter: bara poster med status >= '{args.status_filter}'")

    print("Bygger dokument...")
    doc = Document()

    # Sätt standardfont
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    build_title_page(doc, stats)
    build_toc_section(doc, data)
    build_proficiency_levels_section(doc, data)
    build_competence_section(doc, data, args.status_filter)
    build_glossary_section(doc, data)

    args.output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(args.output)
    print(f"✓ Skapat: {args.output}")


if __name__ == "__main__":
    main()
